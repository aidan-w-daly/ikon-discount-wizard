from docx import Document
import docx
from docx.enum.text import WD_COLOR_INDEX
from docx.shared import Pt
from docx.shared import Inches


IKON_PASS_PRICE = 959
IKON_BASE_PLUS_PASS_PRICE = 889
IKON_BASE_PASS_PRICE = 639

class Student:
    def __init__(self, first_name, last_name, email, ikon_code):
        self.first_name = first_name
        self.last_name = last_name
        self.email = email
        self.ikon_code = ikon_code
    
# def populate_template(Student s) -> str:
#     return f'Hello {s.first_name} {s.last_name},\n\nAs'

#matches google doc exactly BESIDES bullet indent
def generate_promo_code_draft(s: Student):
    document = Document()
    style = document.styles['Normal']
    style.paragraph_format.space_before = Pt(0)
    style.paragraph_format.space_after = Pt(0)
    style.paragraph_format.line_spacing = 1
    font = style.font
    font.name = 'Calibri'
    font.size = Pt(11)
    document.sections[0].left_margin = Inches(1)
    document.sections[0].right_margin = Inches(1)
    document.sections[0].top_margin = Inches(1)
    document.sections[0].bottom_margin = Inches(1)

    p1 = document.add_paragraph(f'Hello {s.first_name} {s.last_name},\n\nAs a current student, you are eligible to receive a special college discount on the Ikon Pass. You can choose one of the following: \n')
    document.add_paragraph(f'Ikon Pass = ${IKON_PASS_PRICE}', style = 'List Bullet')
    document.add_paragraph(f'Ikon Base Plus Pass = ${IKON_BASE_PLUS_PASS_PRICE}', style = 'List Bullet')
    document.add_paragraph(f'Ikon Base Pass = ${IKON_BASE_PASS_PRICE}\n', style = 'List Bullet')
    document.add_paragraph('').add_run(f'Prices are valid through mid-spring and will increase. The official price change date has not been set yet. \n').italic = True
    p2 = document.add_paragraph('Please head to ')
    add_hyperlink(p2, 'https://www.ikonpass.com/', 'IkonPass.com', '0000EE', True)
    p2.add_run(' and select the type of pass you want to purchase. Then proceed to ')
    p2.add_run('VIEW CART').bold = True
    p2.add_run('. Next, on the ')
    p2.add_run('YOUR CART').bold = True
    p2.add_run(' screen, you can select pass insurance here and add the promo code. Click on the ')
    p2.add_run('ADD PROMO CODE').bold = True
    p2.add_run(' box to enter your code below. \n\n')
    p2.add_run('Your promo code is:').bold = True
    p2.add_run(f' {s.ikon_code}\n')
    p2.add_run('This code is registered to your name only. Should it be redeemed by a different person – that pass may be voided pending on student verification.\n\n').italic = True
    p2.add_run('When using the promo code: ')
    document.add_paragraph('\tThe website will automatically enter the code in all capital letters. It is not case sensitive. ', style = 'List Bullet')
    document.add_paragraph('\tAny 0’s in a code are zeros. There are no capital letter O’s in any code.', style = 'List Bullet')
    document.add_paragraph('\tThis is a unique, single use promo code. Once it is used in a transaction, it is no longer active. ', style = 'List Bullet')
    document.add_paragraph('\tThis promo code is registered in your name and cannot be shared with anyone else. ', style = 'List Bullet')
    document.add_paragraph('', style = 'List Bullet').add_run('\tDeadline to use the code is November 30th, 2024.').font.highlight_color = WD_COLOR_INDEX.YELLOW
    document.add_paragraph('\nAfter applying the promo code, please click through the remaining steps to complete your transaction. If you have any Deferral credits from the 23/24 season, those can be applied to your purchase with the promo code for even bigger savings. Any credit will automatically be applied to your purchase when you get to the ASSIGN PASS HOLDER screen.  \n')
    p3 = document.add_paragraph('If you have any issues with your promo code or Ikonpass.com account, you can email ')
    add_hyperlink(p3, 'mailto:college@alterramtnco.com', 'college@alterramtnco.com', '0000EE', True)
    p3.add_run(' and the Ikon Pass College Sales Team can help you.  ')
    document.add_paragraph('\n\nThank you!\n\n\n')

    document.save(f'{s.first_name} {s.last_name}.docx')

# From python-docx github. Why isn't this functionality in the package by default????
def add_hyperlink(paragraph, url, text, color, underline):
    """
    A function that places a hyperlink within a paragraph object.

    :param paragraph: The paragraph we are adding the hyperlink to.
    :param url: A string containing the required url
    :param text: The text displayed for the url
    :return: The hyperlink object
    """

    # This gets access to the document.xml.rels file and gets a new relation id value
    part = paragraph.part
    r_id = part.relate_to(url, docx.opc.constants.RELATIONSHIP_TYPE.HYPERLINK, is_external=True)

    # Create the w:hyperlink tag and add needed values
    hyperlink = docx.oxml.shared.OxmlElement('w:hyperlink')
    hyperlink.set(docx.oxml.shared.qn('r:id'), r_id, )

    # Create a w:r element
    new_run = docx.oxml.shared.OxmlElement('w:r')

    # Create a new w:rPr element
    rPr = docx.oxml.shared.OxmlElement('w:rPr')

    # Add color if it is given
    if not color is None:
      c = docx.oxml.shared.OxmlElement('w:color')
      c.set(docx.oxml.shared.qn('w:val'), color)
      rPr.append(c)

    # Remove underlining if it is requested
    if not underline:
      u = docx.oxml.shared.OxmlElement('w:u')
      u.set(docx.oxml.shared.qn('w:val'), 'none')
      rPr.append(u)

    u = docx.oxml.shared.OxmlElement('w:u')
    u.set(docx.oxml.shared.qn('w:val'), 'single')
    rPr.append(u)

    # Join all the xml elements together add add the required text to the w:r element
    new_run.append(rPr)
    new_run.text = text
    hyperlink.append(new_run)

    paragraph._p.append(hyperlink)

    return hyperlink

    
def send_email(s: Student):
    pass

me = Student('Aidan', 'Daly', 'aidanwdaly@gmail.com', 'x1x1x1x1x1x1')
generate_promo_code_draft(me)